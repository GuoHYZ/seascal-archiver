# -*- coding: utf-8 -*-
"""
DOCX 文档转 TXT 模块。

功能：将 Word 文档（.docx）转换为纯文本文件。
- 段落文本按原始顺序输出
- 表格转换为 Markdown 格式保留结构
- 图片忽略不处理

统一接口：convert(source_path, dest_path) -> int
"""

from pathlib import Path
import traceback

try:
    from docx import Document
except ImportError:
    Document = None


def _escape_md_cell(text):
    """清理表格单元格文本，替换换行和竖线。"""
    return text.replace("\n", " ").replace("\r", " ").replace("|", "\\|")


def _table_to_markdown(table):
    """
    将 python-docx 的 Table 对象转换为 Markdown 表格字符串。
    自动处理合并单元格：空单元格显示为空。
    """
    rows = table.rows
    if not rows:
        return ""

    # 提取所有行数据
    table_data = []
    for row in rows:
        row_data = []
        for cell in row.cells:
            row_data.append(_escape_md_cell(cell.text.strip()))
        table_data.append(row_data)

    if not table_data or not table_data[0]:
        return ""

    # 确定列数（取最大列数，兼容合并单元格导致的列数不一致）
    max_cols = max(len(r) for r in table_data) if table_data else 0
    if max_cols == 0:
        return ""

    # 补齐列数
    for row_data in table_data:
        while len(row_data) < max_cols:
            row_data.append("")

    lines = []

    # 表头行
    header = table_data[0]
    lines.append("| " + " | ".join(header) + " |")

    # 分隔行
    lines.append("|" + "|".join([" --- " for _ in range(max_cols)]) + "|")

    # 数据行
    for row_data in table_data[1:]:
        lines.append("| " + " | ".join(row_data) + " |")

    return "\n".join(lines)


def _extract_docx_content(doc):
    """
    按文档原始顺序遍历段落和表格，返回文本行列表。
    文档结构：paragraph 和 table 按顺序交替出现。
    """
    lines = []

    # python-docx 不直接提供顺序迭代器，通过 iterblock_items 模式处理
    # 方式：遍历 document.body 的 XML 子元素，判断是段落还是表格
    from docx.oxml.ns import qn

    body = doc.element.body

    for child in body:
        tag = child.tag

        # 段落
        if tag == qn("w:p"):
            text = ""
            try:
                # 获取该元素对应的 Paragraph 对象
                # 通过 Element 定位到对应的 paragraph
                for para in doc.paragraphs:
                    if para._element is child:
                        text = para.text.strip()
                        break
            except Exception:
                pass

            if text:
                lines.append(text)

        # 表格
        elif tag == qn("w:tbl"):
            try:
                for table in doc.tables:
                    if table._element is child:
                        md_table = _table_to_markdown(table)
                        if md_table:
                            lines.append("")
                            lines.append(md_table)
                            lines.append("")
                        break
            except Exception:
                pass

    return lines


def docx_to_txt(source_path, dest_path):
    """
    将 DOCX 文件转换为 TXT 文件。

    source_path: str 或 Path — 源文件路径
    dest_path:   str 或 Path — 目标 TXT 文件路径
    """
    if Document is None:
        raise ImportError("请安装 python-docx: pip install python-docx")

    source_path = Path(source_path)
    dest_path = Path(dest_path)

    doc = Document(str(source_path))

    lines = _extract_docx_content(doc)

    dest_path.parent.mkdir(parents=True, exist_ok=True)

    with open(dest_path, "w", encoding="utf-8", errors="replace") as f:
        f.write(f"文件：{source_path.name}\n")
        f.write("=" * 60 + "\n\n")

        for line in lines:
            f.write(line + "\n")

        if not lines:
            f.write("(文档无可提取文本)\n")


def convert(source_path, dest_path):
    """
    统一转换接口。

    source_path: str 或 Path — 源 DOCX 文件路径
    dest_path:   str 或 Path — 目标 TXT 文件路径
    返回: int — 提取的总字符数
    """
    docx_to_txt(source_path, dest_path)

    dest_path = Path(dest_path)
    if dest_path.exists():
        return len(dest_path.read_text(encoding="utf-8", errors="replace"))
    return 0


# ============================================================
# 独立运行：批量转换当前目录下指定文件夹的 DOCX 文件
# ============================================================
if __name__ == "__main__":
    input_dir = Path("./DOCX")
    output_dir = Path("./TXT")

    if not input_dir.is_dir():
        print(f"[错误] 输入目录不存在: {input_dir}")
    else:
        output_dir.mkdir(exist_ok=True)
        docx_files = list(input_dir.glob("*.docx"))
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]

        if not docx_files:
            print(f"[警告] 目录 {input_dir} 中未找到任何 DOCX 文件")
        else:
            success, failed = 0, 0
            total = len(docx_files)
            for i, docx in enumerate(docx_files, start=1):
                txt = output_dir / (docx.stem + ".txt")
                try:
                    char_count = convert(docx, txt)
                    success += 1
                    print(f"[{i}/{total}] [OK] {docx.name} ({char_count} 字符)")
                except Exception as e:
                    failed += 1
                    print(f"[{i}/{total}] [FAIL] {docx.name}")
                    print(f"       原因: {e}")
                    traceback.print_exc()
            print(f"\n处理结束 - 共 {total} 个文件")
            print(f"成功: {success}, 失败: {failed}")